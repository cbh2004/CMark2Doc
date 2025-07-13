#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
简单有效的数学公式转换器
使用Unicode数学符号和特殊格式来模拟数学公式效果
"""

import re
import markdown
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup

class SimpleEffectiveConverter:
    def __init__(self):
        self.doc = Document()
        self.setup_document_styles()
    
    def setup_document_styles(self):
        """设置文档样式"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
    
    def latex_to_unicode_math(self, latex_code):
        """将LaTeX转换为Unicode数学表示"""
        latex_code = latex_code.strip()
        
        # 处理分数 - 使用分数线字符
        latex_code = self.process_fractions(latex_code)
        
        # 处理上下标
        latex_code = self.process_scripts(latex_code)
        
        # 处理数学符号
        latex_code = self.process_symbols(latex_code)
        
        return latex_code
    
    def process_fractions(self, text):
        """处理分数，使用Unicode分数表示"""
        # 处理简单分数
        simple_fractions = {
            '\\frac{1}{2}': '½',
            '\\frac{1}{3}': '⅓',
            '\\frac{2}{3}': '⅔',
            '\\frac{1}{4}': '¼',
            '\\frac{3}{4}': '¾',
            '\\frac{1}{5}': '⅕',
            '\\frac{2}{5}': '⅖',
            '\\frac{3}{5}': '⅗',
            '\\frac{4}{5}': '⅘',
            '\\frac{1}{6}': '⅙',
            '\\frac{5}{6}': '⅚',
            '\\frac{1}{8}': '⅛',
            '\\frac{3}{8}': '⅜',
            '\\frac{5}{8}': '⅝',
            '\\frac{7}{8}': '⅞'
        }
        
        for latex_frac, unicode_frac in simple_fractions.items():
            text = text.replace(latex_frac, unicode_frac)
        
        # 处理复杂分数 - 使用分数线表示法
        def replace_complex_frac(match):
            num = match.group(1)
            den = match.group(2)
            
            # 递归处理分子分母
            num = self.process_nested_expression(num)
            den = self.process_nested_expression(den)
            
            # 使用分数线格式
            return f"({num})/({den})"
        
        # 匹配复杂分数
        text = re.sub(r'\\frac\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}', 
                     replace_complex_frac, text)
        
        return text
    
    def process_nested_expression(self, expr):
        """处理嵌套表达式"""
        # 处理上标
        expr = self.process_scripts(expr)
        
        # 处理符号
        expr = self.process_symbols(expr)
        
        # 处理嵌套分数
        if '\\frac' in expr:
            expr = self.process_fractions(expr)
        
        return expr
    
    def process_scripts(self, text):
        """处理上下标"""
        # 处理上下标组合 A_1^2 -> A₁²
        def replace_subsup(match):
            base = match.group(1)
            sub = match.group(2)
            sup = match.group(3)
            return base + self.to_subscript(sub) + self.to_superscript(sup)
        
        text = re.sub(r'([A-Za-z])_\{?([^}]+)\}?\^\{?([^}]+)\}?', replace_subsup, text)
        
        # 处理下标 B_{12} -> B₁₂
        def replace_sub(match):
            base = match.group(1)
            sub = match.group(2)
            return base + self.to_subscript(sub)
        
        text = re.sub(r'([A-Za-z])_\{?([^}]+)\}?', replace_sub, text)
        
        # 处理上标 x^2 -> x²
        def replace_sup(match):
            base = match.group(1)
            sup = match.group(2)
            return base + self.to_superscript(sup)
        
        text = re.sub(r'([A-Za-z])\^\{?([^}]+)\}?', replace_sup, text)
        
        return text
    
    def to_subscript(self, text):
        """转换为下标Unicode字符"""
        subscript_map = {
            '0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄',
            '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉',
            '+': '₊', '-': '₋', '=': '₌', '(': '₍', ')': '₎',
            'a': 'ₐ', 'e': 'ₑ', 'i': 'ᵢ', 'o': 'ₒ', 'u': 'ᵤ',
            'x': 'ₓ'
        }
        result = ''
        for char in text:
            result += subscript_map.get(char, char)
        return result
    
    def to_superscript(self, text):
        """转换为上标Unicode字符"""
        superscript_map = {
            '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
            '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
            '+': '⁺', '-': '⁻', '=': '⁼', '(': '⁽', ')': '⁾',
            'a': 'ᵃ', 'b': 'ᵇ', 'c': 'ᶜ', 'd': 'ᵈ', 'e': 'ᵉ',
            'f': 'ᶠ', 'g': 'ᵍ', 'h': 'ʰ', 'i': 'ⁱ', 'j': 'ʲ',
            'k': 'ᵏ', 'l': 'ˡ', 'm': 'ᵐ', 'n': 'ⁿ', 'o': 'ᵒ',
            'p': 'ᵖ', 'r': 'ʳ', 's': 'ˢ', 't': 'ᵗ', 'u': 'ᵘ',
            'v': 'ᵛ', 'w': 'ʷ', 'x': 'ˣ', 'y': 'ʸ', 'z': 'ᶻ'
        }
        result = ''
        for char in text:
            result += superscript_map.get(char, char)
        return result
    
    def process_symbols(self, text):
        """处理数学符号"""
        # 希腊字母
        greek_letters = {
            '\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ', '\\delta': 'δ',
            '\\epsilon': 'ε', '\\varepsilon': 'ε', '\\zeta': 'ζ', '\\eta': 'η',
            '\\theta': 'θ', '\\vartheta': 'θ', '\\iota': 'ι', '\\kappa': 'κ',
            '\\lambda': 'λ', '\\mu': 'μ', '\\nu': 'ν', '\\xi': 'ξ',
            '\\pi': 'π', '\\varpi': 'π', '\\rho': 'ρ', '\\varrho': 'ρ',
            '\\sigma': 'σ', '\\varsigma': 'ς', '\\tau': 'τ', '\\upsilon': 'υ',
            '\\phi': 'φ', '\\varphi': 'φ', '\\chi': 'χ', '\\psi': 'ψ', '\\omega': 'ω',
            '\\Gamma': 'Γ', '\\Delta': 'Δ', '\\Theta': 'Θ', '\\Lambda': 'Λ',
            '\\Xi': 'Ξ', '\\Pi': 'Π', '\\Sigma': 'Σ', '\\Upsilon': 'Υ',
            '\\Phi': 'Φ', '\\Psi': 'Ψ', '\\Omega': 'Ω'
        }
        
        for latex_symbol, unicode_symbol in greek_letters.items():
            text = text.replace(latex_symbol, unicode_symbol)
        
        # 运算符
        operators = {
            '\\pm': '±', '\\mp': '∓', '\\times': '×', '\\div': '÷',
            '\\cdot': '·', '\\ast': '∗', '\\star': '⋆', '\\circ': '∘',
            '\\bullet': '•', '\\cap': '∩', '\\cup': '∪',
            '\\leq': '≤', '\\le': '≤', '\\geq': '≥', '\\ge': '≥',
            '\\neq': '≠', '\\ne': '≠', '\\equiv': '≡', '\\sim': '∼',
            '\\approx': '≈', '\\cong': '≅', '\\propto': '∝',
            '\\in': '∈', '\\notin': '∉', '\\subset': '⊂', '\\supset': '⊃',
            '\\subseteq': '⊆', '\\supseteq': '⊇', '\\emptyset': '∅',
            '\\infty': '∞', '\\partial': '∂', '\\nabla': '∇',
            '\\forall': '∀', '\\exists': '∃', '\\neg': '¬'
        }
        
        for latex_symbol, unicode_symbol in operators.items():
            text = text.replace(latex_symbol, unicode_symbol)
        
        # 积分和求和
        text = text.replace('\\int', '∫')
        text = text.replace('\\iint', '∬')
        text = text.replace('\\iiint', '∭')
        text = text.replace('\\oint', '∮')
        text = text.replace('\\sum', 'Σ')
        text = text.replace('\\prod', 'Π')
        
        # 箭头
        arrows = {
            '\\rightarrow': '→', '\\to': '→', '\\leftarrow': '←', '\\gets': '←',
            '\\leftrightarrow': '↔', '\\Rightarrow': '⇒', '\\Leftarrow': '⇐',
            '\\Leftrightarrow': '⇔', '\\mapsto': '↦', '\\uparrow': '↑',
            '\\downarrow': '↓', '\\updownarrow': '↕'
        }
        
        for latex_symbol, unicode_symbol in arrows.items():
            text = text.replace(latex_symbol, unicode_symbol)
        
        # 函数名
        functions = {
            '\\sin': 'sin', '\\cos': 'cos', '\\tan': 'tan', '\\cot': 'cot',
            '\\sec': 'sec', '\\csc': 'csc', '\\log': 'log', '\\ln': 'ln',
            '\\exp': 'exp', '\\max': 'max', '\\min': 'min', '\\lim': 'lim'
        }
        
        for latex_symbol, unicode_symbol in functions.items():
            text = text.replace(latex_symbol, unicode_symbol)
        
        # 清理剩余的LaTeX命令
        text = re.sub(r'\\[a-zA-Z]+', '', text)
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def create_equation_paragraph(self, latex_code):
        """创建方程式段落"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 处理多行公式
        if '\\\\' in latex_code:
            expressions = [expr.strip() for expr in latex_code.split('\\\\') if expr.strip()]
            
            for i, expr in enumerate(expressions):
                if i > 0:
                    # 添加空格分隔
                    p.add_run("     ")
                
                # 转换为Unicode数学表示
                math_text = self.latex_to_unicode_math(expr)
                run = p.add_run(math_text)
                run.font.name = 'Cambria Math'
                run.font.size = Pt(16)
        else:
            # 转换为Unicode数学表示
            math_text = self.latex_to_unicode_math(latex_code)
            run = p.add_run(math_text)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(16)
    
    def add_inline_equation(self, paragraph, latex_code):
        """添加行内方程式"""
        # 转换为Unicode数学表示
        math_text = self.latex_to_unicode_math(latex_code)
        run = paragraph.add_run(math_text)
        run.font.name = 'Cambria Math'
        run.font.size = Pt(12)
    
    def convert_markdown_content(self, markdown_content):
        """转换Markdown内容"""
        # 提取数学公式
        math_formulas = []
        formula_counter = 0
        
        # 处理块级公式
        def replace_block_math(match):
            nonlocal formula_counter
            formula_id = f"BLOCK_MATH_{formula_counter}"
            math_formulas.append({
                'id': formula_id,
                'content': match.group(1),
                'type': 'block'
            })
            formula_counter += 1
            return f"\n{formula_id}\n"
        
        processed_content = re.sub(r'\$\$([^$]+?)\$\$', replace_block_math, markdown_content, flags=re.DOTALL)
        
        # 处理行内公式
        def replace_inline_math(match):
            nonlocal formula_counter
            formula_id = f"INLINE_MATH_{formula_counter}"
            math_formulas.append({
                'id': formula_id,
                'content': match.group(1),
                'type': 'inline'
            })
            formula_counter += 1
            return formula_id
        
        processed_content = re.sub(r'\$([^$\n]+?)\$', replace_inline_math, processed_content)
        
        # 转换Markdown为HTML
        md = markdown.Markdown(extensions=['extra', 'codehilite', 'toc', 'tables'])
        html_content = md.convert(processed_content)
        
        # 解析HTML并转换为Word
        soup = BeautifulSoup(html_content, 'html.parser')
        self.convert_html_to_word(soup, math_formulas)

    def convert_html_to_word(self, soup, math_formulas):
        """将HTML转换为Word文档"""
        formula_map = {formula['id']: formula for formula in math_formulas}

        for element in soup.find_all(recursive=False):
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                text = element.get_text()
                text = self.process_text_with_formulas(text, formula_map)
                self.doc.add_heading(text, level=level)

            elif element.name == 'p':
                text = element.get_text()
                self.process_paragraph_with_formulas(text, formula_map)

            elif element.name == 'ul':
                # 处理无序列表
                for li in element.find_all('li'):
                    text = li.get_text()
                    p = self.doc.add_paragraph(style='List Bullet')
                    self.process_paragraph_text_with_formulas(p, text, formula_map)

    def process_paragraph_with_formulas(self, text, formula_map):
        """处理包含公式的段落"""
        formula_pattern = r'(BLOCK_MATH_\d+|INLINE_MATH_\d+)'
        parts = re.split(formula_pattern, text)

        if len(parts) == 1:
            if text.strip():
                self.doc.add_paragraph(text)
        else:
            current_paragraph = None

            for part in parts:
                if part in formula_map:
                    formula = formula_map[part]
                    if formula['type'] == 'block':
                        if current_paragraph is not None:
                            current_paragraph = None
                        self.create_equation_paragraph(formula['content'])
                    else:
                        if current_paragraph is None:
                            current_paragraph = self.doc.add_paragraph()
                        self.add_inline_equation(current_paragraph, formula['content'])
                elif part.strip():
                    if current_paragraph is None:
                        current_paragraph = self.doc.add_paragraph()
                    current_paragraph.add_run(part)

    def process_paragraph_text_with_formulas(self, paragraph, text, formula_map):
        """在现有段落中处理包含公式的文本"""
        formula_pattern = r'(BLOCK_MATH_\d+|INLINE_MATH_\d+)'
        parts = re.split(formula_pattern, text)

        for part in parts:
            if part in formula_map:
                formula = formula_map[part]
                if formula['type'] == 'inline':
                    self.add_inline_equation(paragraph, formula['content'])
            elif part.strip():
                paragraph.add_run(part)

    def process_text_with_formulas(self, text, formula_map):
        """处理包含公式的文本（用于标题）"""
        formula_pattern = r'(BLOCK_MATH_\d+|INLINE_MATH_\d+)'
        parts = re.split(formula_pattern, text)

        result = ""
        for part in parts:
            if part in formula_map:
                formula = formula_map[part]
                # 对于标题，使用简化的表示
                result += self.latex_to_unicode_math(formula['content'])
            else:
                result += part

        return result
